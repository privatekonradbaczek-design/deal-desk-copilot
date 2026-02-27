from __future__ import annotations

import hashlib
import re
from uuid import UUID

import structlog
import tiktoken

from indexing_service.domain.models import DocumentChunk, IndexedChunk

logger = structlog.get_logger(__name__)


class ChunkingService:
    """Token-aware text chunking with configurable size and overlap."""

    def __init__(
        self,
        chunk_size_tokens: int = 512,
        overlap_tokens: int = 64,
        encoding_name: str = "cl100k_base",
    ) -> None:
        self._chunk_size = chunk_size_tokens
        self._overlap = overlap_tokens
        self._enc = tiktoken.get_encoding(encoding_name)

    def chunk_text(
        self,
        document_id: UUID,
        tenant_id: str,
        text: str,
        page_number: int | None = None,
    ) -> list[DocumentChunk]:
        tokens = self._enc.encode(text)
        chunks: list[DocumentChunk] = []
        step = self._chunk_size - self._overlap
        chunk_index = 0

        for start in range(0, len(tokens), step):
            chunk_tokens = tokens[start : start + self._chunk_size]
            if not chunk_tokens:
                break
            content = self._enc.decode(chunk_tokens)
            chunks.append(
                DocumentChunk(
                    document_id=document_id,
                    tenant_id=tenant_id,
                    content=content,
                    page_number=page_number,
                    chunk_index=chunk_index,
                    token_count=len(chunk_tokens),
                )
            )
            chunk_index += 1

        logger.debug(
            "chunking.completed",
            document_id=str(document_id),
            token_count=len(tokens),
            chunk_count=len(chunks),
        )
        return chunks


class TextExtractionService:
    """Extracts plain text from supported document formats."""

    def extract_from_pdf(self, content: bytes) -> tuple[str, int]:
        from pypdf import PdfReader
        import io

        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages), len(reader.pages)

    def extract_from_docx(self, content: bytes) -> tuple[str, int]:
        from docx import Document
        import io

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs), len(doc.paragraphs)

    def extract_from_text(self, content: bytes) -> tuple[str, int]:
        text = content.decode("utf-8", errors="replace")
        page_count = max(1, len(text) // 3000)
        return text, page_count

    def extract(self, content: bytes, content_type: str) -> tuple[str, int]:
        if "pdf" in content_type:
            return self.extract_from_pdf(content)
        elif "wordprocessingml" in content_type:
            return self.extract_from_docx(content)
        else:
            return self.extract_from_text(content)
